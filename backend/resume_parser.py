from io import BytesIO

from pypdf import PdfReader
from docx import Document


class UnsupportedResumeFormat(Exception):
    pass


def extract_text(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")
    raise UnsupportedResumeFormat(f"Unsupported file type: {filename}")


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages).strip()


def _extract_docx(data: bytes) -> str:
    doc = Document(BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs).strip()
